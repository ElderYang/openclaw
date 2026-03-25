#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
创建 AI 科技日报 Word 文档
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
from datetime import datetime

# 创建文档
doc = Document()

# 设置标题
title = doc.add_heading('📰 AI 科技日报', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 日期
date = datetime.now().strftime('%Y 年 %m 月 %d 日 %A')
date_para = doc.add_paragraph(date)
date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
date_para.runs[0].italic = True

doc.add_paragraph()

# 内容
content = """
🔥 头条要闻

OpenAI 与 Google 同日发布新模型
• OpenAI：3 月 4 日正式推送 GPT-5.3 Instant，覆盖全量 ChatGPT 用户
• Google：同日发布 Gemini 3.1 Flash-Lite，号称 Gemini 3 系列"速度最快、最具成本效益"的模型
• 两大 AI 巨头同一天各甩出一张牌，市场竞争进一步加剧

---

🌍 国际大厂动态

OpenAI
• 支出计划：预计 2026 年支出达 170 亿美元，2027 年 350 亿美元，2028 年 450 亿美元
• Azure 合作：承诺在 Azure 上额外支出 2500 亿美元
• GPT-5.4：传闻凌晨发布新版本（待确认）

Google
• Gemini 系列：Gemini 3.1 Flash-Lite 正式发布，主打速度和成本效益
• 数据中心：签署数据中心电力支付承诺（但亚利桑那州 3 个项目未包含）

Meta & 微软
• 微软：持有 OpenAI 27% 股份，Azure 成为主要 AI 基础设施提供商
• Anthropic：承诺在 Azure 上支出 300 亿美元
• 科技联盟：Google、Meta、微软、OpenAI、xAI、Oracle、AWS 签署数据中心电力协议

NVIDIA
• GTC 2026：黄仁勋将主持全球科技领袖会议，展示"AI 时代"最新成果
• 参会企业：Google DeepMind、Meta、微软、OpenAI、特斯拉等

---

🇨🇳 国内大厂动态

阿里
• 通义千问：API 价格持续调整，保持市场竞争力
• 阿里云开发者社区：发布 2026 企业建站新逻辑指南，应对 AI 时代搜索引擎算法变化

字节跳动
• 豆包：持续优化用户体验，早期发展策略引发关注
• AI 硬件：豆包 AI 智能体耳机 Ola Friend 已推向市场

DeepSeek
• 开源战略：代表开源大模型阵营，与闭源模型形成"双核驱动"
• 本地部署：支持在 Mac 集群上运行万亿参数模型（4 台 M3 Ultra 可组建 2TB 显存池）

Kimi（月之暗面）
• API 价格：参与国内大模型价格竞争
• 模型能力：支持在本地 Mac 集群流畅运行

腾讯
• 混元大模型：API 价格参与市场竞争

MiniMax
• Agent 平台：推出全能智能体工具，支持一键云部署
• 云端服务：上线云端 Agent 服务

---

💡 技术突破

AI 智能体（Agent）
• 自主进化：工程师开发出首个能自主进化的 AI 系统
• Cursor：上线程序员专用 AI 助手 Cursor Cloud Agents
• Perplexity：推出云端 AI 助手 Perplexity Computer
• OpenClaw：预置专家模式，告别命令行

AI 硬件
• 高通：推出新款 AI 可穿戴设备芯片
• AI 眼镜：2026 年 AI 智能眼镜成为新热点
• Motorola Razr 2026：折叠屏手机搭载 18GB RAM + 1TB 存储

本地部署
• llmfit 工具：1 秒检测电脑能运行哪些 AI 大模型（支持 500+ 模型）
• Mac 集群：4 台 M3 Ultra 可组建个人 AI 算力中心，运行万亿参数模型

---

📊 行业趋势

2026 AI 创新机会
• 双核驱动：闭源（OpenAI 主导）与开源（DeepSeek 主导）并行发展
• 地缘竞争：美国与中国成为 AI 发展两大核心力量
• 应用场景：AI 向垂直行业深度渗透，研发周期数量级缩短

搜索引擎变革
• 算法迭代：2025 年底搜索引擎算法重大更新
• AI 抓取：大模型抓取规则全面落地
• 建站逻辑：传统企业网站逻辑已过时，需适应 AI 时代流量分发

手机 AI 落地
• 国产手机：华为、荣耀、小米、VIVO、OPPO 纷纷搭载端侧大模型
• 功能应用：AI 消除、AI 搜索对话、生活助手等已提供较好体验
• 智能体：荣耀 Magic7 系列 YOYO 智能体可实现一句话点咖啡等功能

---

🎯 产业活动

智能向善倡议
• 发起方：中国企业家木兰汇（《中国企业家》杂志社）
• 参与者：43 位女性企业家领衔
• 目标：推动 AI 朝着有益、安全、公平的方向发展
• 关注点：技术迭代与伦理治理、效率与公平、数据规模与隐私保护

2026 AI 最佳场景渗透案例征集
• 主办方：36 氪
• 截止时间：倒计时 10 天
• 评判标准：
  1. 研发与验证周期能否实现数量级缩短
  2. 人机协同深度与整体效率是否显著提升
  3. 最终产品或服务质量是否获得实质性飞跃

---

📈 投资观察

热门赛道
• AI 视频：DeepSeek + ComfyUI + SD 组合，3 分钟生成爆款 AI 视频
• AI 短剧：日入 5 位数的 AI 视频赛道
• AI Agent：企业级智能体搭建需求旺盛

关注股票
• 微软：持有 OpenAI 股份，Azure 受益于 AI 支出增长
• NVIDIA：GTC 2026 即将召开，AI 芯片需求持续

---

🔗 今日推荐

• llmfit：检测你的电脑能跑哪些 AI 大模型（小众软件推荐）
• Gemini Voyager 插件：让 Gemini 使用体验翻倍（GitHub: github.com/Nagi-ovo/gemini-voyager）
• B 站教程：《20 秒部署 AI 助手！MaxClaw 零成本玩转 AI 短剧、PPT、炒股分析》

---

数据来源：36 氪、新浪财经、B 站、GitHub 等公开信息
更新时间：2026 年 3 月 9 日 15:00（Asia/Shanghai）

AI 科技日报 · 每日更新
"""

doc.add_paragraph(content)

# 保存文档
desktop_path = os.path.join(os.path.expanduser("~"), "Desktop")
doc_path = os.path.join(desktop_path, "AI 科技日报_20260309.docx")
doc.save(doc_path)

print(f"文档已创建：{doc_path}")
